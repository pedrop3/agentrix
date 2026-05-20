"""
Extracao de texto a partir de arquivos enviados pelo usuario.

Suportado:
  - PDF  (pypdf)
  - DOCX (python-docx)
  - TXT / MD (decodificacao direta)
  - Imagens (sem OCR aqui; salva o arquivo mas nao indexa texto)
"""
from __future__ import annotations

import io
from dataclasses import dataclass
from typing import Optional

PDF_EXTS = {".pdf"}
DOCX_EXTS = {".docx"}
TEXT_EXTS = {".txt", ".md", ".markdown", ".rst", ".log", ".csv"}
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".heic", ".bmp"}


@dataclass
class Extracted:
    text: str
    kind: str  # "pdf" | "docx" | "text" | "image" | "unknown"
    note: Optional[str] = None  # mensagem extra quando nao indexavel


def detect_kind(filename: str, content_type: Optional[str] = None) -> str:
    name = (filename or "").lower()
    for ext in PDF_EXTS:
        if name.endswith(ext):
            return "pdf"
    for ext in DOCX_EXTS:
        if name.endswith(ext):
            return "docx"
    for ext in TEXT_EXTS:
        if name.endswith(ext):
            return "text"
    for ext in IMAGE_EXTS:
        if name.endswith(ext):
            return "image"
    # fallback por content-type
    ct = (content_type or "").lower()
    if "pdf" in ct:
        return "pdf"
    if "word" in ct or "officedocument" in ct:
        return "docx"
    if ct.startswith("text/") or "markdown" in ct:
        return "text"
    if ct.startswith("image/"):
        return "image"
    return "unknown"


def extract(data: bytes, filename: str, content_type: Optional[str] = None) -> Extracted:
    kind = detect_kind(filename, content_type)
    if kind == "pdf":
        return Extracted(text=_extract_pdf(data), kind="pdf")
    if kind == "docx":
        return Extracted(text=_extract_docx(data), kind="docx")
    if kind == "text":
        return Extracted(text=_extract_text(data), kind="text")
    if kind == "image":
        return Extracted(
            text="",
            kind="image",
            note="Imagem recebida; o modelo atual nao processa imagens (sem OCR). Arquivo salvo apenas.",
        )
    return Extracted(
        text="",
        kind="unknown",
        note=f"Tipo nao suportado: {filename}",
    )


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n\n".join(p for p in parts if p)


def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    # tabelas
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_text(data: bytes) -> str:
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")
